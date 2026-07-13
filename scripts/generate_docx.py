#!/usr/bin/env python3
"""Combine all .md files in a directory into a single formatted Word document."""
import sys, re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_cn_font(run, name='SimSun', size=Pt(11)):
    run.font.name = name
    run.font.size = size
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {})
    rFonts.set(qn('w:eastAsia'), name)
    rPr.insert(0, rFonts)

def add_title(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        set_cn_font(run, 'SimHei')
    else:
        run = p.add_run(text)
        run.font.size = [Pt(16), Pt(14), Pt(12)][min(level-1, 2)]
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        set_cn_font(run, 'SimHei')
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

def add_body(doc, text):
    if not text.strip():
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.font.bold = True
        else:
            run = p.add_run(part)
        set_cn_font(run, 'SimSun', Pt(11))

def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    shd = p._element.get_or_add_pPr()
    s = shd.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): 'F5F5F5'
    })
    shd.append(s)

def process_md_file(doc, filepath):
    filename = Path(filepath).stem
    add_title(doc, filename, level=1)
    doc.add_paragraph()

    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    in_code = False
    code_buf = []

    for line in lines:
        if line.strip().startswith('```'):
            if in_code:
                add_code_block(doc, '\n'.join(code_buf))
                code_buf = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_buf.append(line)
            continue

        if line.startswith('# '):
            add_title(doc, line[2:], level=0)
        elif line.startswith('## '):
            add_title(doc, line[3:], level=1)
        elif line.startswith('### '):
            add_title(doc, line[4:], level=2)
        elif line.startswith('#### '):
            add_title(doc, line[5:], level=3)
        elif line.strip().startswith('|') and line.strip().endswith('|'):
            if '---' in line:
                continue
            cells = [c.strip() for c in line.strip().split('|')[1:-1]]
            add_body(doc, '  |  '.join(cells))
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            add_body(doc, '  •  ' + line.strip()[2:])
        elif re.match(r'^\d+\.\s', line.strip()):
            add_body(doc, '  ' + line.strip())
        elif line.strip() == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run('─' * 60)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            set_cn_font(run, 'SimSun', Pt(6))
        elif line.strip():
            add_body(doc, line)

def main():
    if len(sys.argv) != 3:
        print(f"Usage: {sys.argv[0]} <md_directory> <output.docx>")
        sys.exit(1)

    md_dir = sys.argv[1]
    docx_path = sys.argv[2]

    # Collect and sort .md files
    md_files = sorted(
        [f for f in Path(md_dir).glob('*.md')
         if f.name != '.gitkeep'],
        key=lambda f: f.name
    )

    if not md_files:
        print("No .md files found")
        sys.exit(1)

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style.font.size = Pt(11)

    # Cover page
    dir_name = Path(md_dir).name
    add_title(doc, dir_name, level=0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{len(md_files)} 个内容文件')
    set_cn_font(run, 'SimSun', Pt(10))
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # Process each .md file
    for md_file in md_files:
        process_md_file(doc, str(md_file))
        doc.add_paragraph()

    doc.save(docx_path)
    print(f"✅ {docx_path} ({len(md_files)} 个文件)")

if __name__ == '__main__':
    main()
